"""
Text extraction from PDF and DOCX resume files.
"""
import io
import pdfplumber
from docx import Document


def extract_text(content: bytes, filename: str) -> str:
    """Extract plain text from a PDF or DOCX file."""
    if filename.lower().endswith(".pdf"):
        return _extract_from_pdf(content)
    elif filename.lower().endswith(".docx"):
        return _extract_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def _extract_from_pdf(content: bytes) -> str:
    """Extract text from a PDF using pdfplumber."""
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text.strip())
    return "\n\n".join(text_parts)


def _extract_from_docx(content: bytes) -> str:
    """Extract text from a DOCX using python-docx."""
    doc = Document(io.BytesIO(content))
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)
    return "\n".join(paragraphs)
