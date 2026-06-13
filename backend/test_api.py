import requests
from docx import Document
import io, time

doc = Document()
doc.add_heading('Chinmay Sharma', 0)
doc.add_paragraph('Full Stack Developer | Python | Machine Learning | REST APIs | SQL | Git | Flask | TensorFlow')
doc.add_heading('Experience', level=1)
doc.add_paragraph('Worked on backend development for the company project using Python')
doc.add_paragraph('Helped with machine learning model for data classification')
doc.add_paragraph('Made improvements to the existing codebase and fixed bugs')
doc.add_paragraph('Built REST API endpoints using Flask framework')
doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, Flask, TensorFlow, SQL, Git, NumPy, Pandas, Data Analysis, REST API')

buf = io.BytesIO()
doc.save(buf)
buf.seek(0)

jd = """Senior Backend Engineer
We are looking for expertise in:
Python, Docker, Kubernetes, AWS, CI/CD pipelines, Microservices, System Design,
Agile/Scrum, Redis caching, GraphQL, REST APIs, PostgreSQL, MongoDB."""

print('Sending to Claude API...')
start = time.time()

resp = requests.post(
    'http://localhost:8000/analyze',
    files={'file': ('resume.docx', buf, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
    data={'job_description': jd},
    timeout=60
)

elapsed = round(time.time() - start, 1)
data = resp.json()
print(f'Response in {elapsed}s | Status: {resp.status_code}')
print(f'Is Demo: {data["is_demo"]}')
print(f'ATS Score: {data["ats_score"]}')
print(f'Summary: {data["summary"][:120]}...')
print(f'Matched ({len(data["matched_keywords"])}): {data["matched_keywords"]}')
print(f'Missing ({len(data["missing_keywords"])}): {data["missing_keywords"]}')
print(f'Weak bullets: {len(data["weak_bullets"])}')
if not data['is_demo']:
    print('\nSUCCESS - Live Claude analysis is working!')
else:
    print('\nWARNING - Still returning demo data')
